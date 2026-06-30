from __future__ import annotations

import os
import tempfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Bao_cao_toan_van_rut_gon_6_trang_giu_format_bang_landscape_tap_chi_ncv_app.docx"
OUT = ROOT / os.getenv(
    "REHAB_OUT_DOCX",
    "Bao_cao_toan_van_rut_gon_6_trang_giu_format_bang_landscape_tap_chi_ncv_app_4anh_moi_trang.docx",
)


def set_run_font(run, size: float = 7.2, *, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.italic = italic


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def fitted_size(path: Path, max_width_cm: float, max_height_cm: float) -> tuple[float, float]:
    with Image.open(path) as img:
        width_px, height_px = img.size
    width_cm = max_width_cm
    height_cm = width_cm * height_px / width_px
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * width_px / height_px
    return width_cm, height_cm


def figure_tables(doc: Document):
    tables = []
    for table in doc.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "Hình PL4." in text:
            tables.append(table)
    return tables


def extract_figures(doc: Document, tmpdir: Path) -> list[tuple[Path, str]]:
    figures: list[tuple[Path, str]] = []
    index = 1
    for table in figure_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                captions = [p.text.strip() for p in cell.paragraphs if p.text.strip().startswith("Hình PL4.")]
                if not captions:
                    continue
                blip = next(cell._tc.iter(qn("a:blip")), None)
                if blip is None:
                    continue
                rid = blip.get(qn("r:embed"))
                if not rid:
                    continue
                part = doc.part.related_parts[rid]
                ext = Path(str(part.partname)).suffix or ".png"
                image_path = tmpdir / f"pl4_{index:02d}{ext}"
                image_path.write_bytes(part.blob)
                figures.append((image_path, captions[-1]))
                index += 1
    return figures


def remove_old_figure_area(doc: Document) -> None:
    tables = figure_tables(doc)
    if not tables:
        raise RuntimeError("Cannot find PL4 figure tables.")
    body = doc.element.body
    children = list(body)
    start = children.index(tables[0]._element)
    for child in children[start:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def add_figure_grid(doc: Document, figures: list[tuple[Path, str]]) -> None:
    for group_start in range(0, len(figures), 4):
        if group_start:
            doc.add_page_break()
        group = figures[group_start : group_start + 4]
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        remove_table_borders(table)
        for cell in table._cells:
            set_cell_width(cell, 12.0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
        for idx, (image_path, caption) in enumerate(group):
            row = idx // 2
            col = idx % 2
            cell = table.rows[row].cells[col]
            image_p = cell.paragraphs[0]
            image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_p.paragraph_format.space_after = Pt(1)
            width_cm, height_cm = fitted_size(image_path, 11.2, 5.95)
            image_p.add_run().add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))
            cap = cell.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.line_spacing = 1
            cap.paragraph_format.space_before = Pt(0)
            cap.paragraph_format.space_after = Pt(1)
            run = cap.add_run(caption)
            set_run_font(run, size=7.2, italic=True)


def main() -> None:
    doc = Document(str(SRC))
    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        figures = extract_figures(doc, tmpdir)
        if len(figures) < 4:
            raise RuntimeError(f"Expected PL4 figures, found {len(figures)}.")
        remove_old_figure_area(doc)
        add_figure_grid(doc, figures)
        doc.save(OUT)
    print(OUT)
    print(f"figures={len(figures)} groups={(len(figures) + 3) // 4}")


if __name__ == "__main__":
    main()
