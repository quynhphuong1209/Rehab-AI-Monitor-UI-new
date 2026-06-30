from __future__ import annotations

import os
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Bao_cao_toan_van_rut_gon_6_trang_giu_format_bang_landscape_tap_chi_ncv_app_4anh_moi_trang.docx"
OUT = ROOT / os.getenv(
    "REHAB_OUT_DOCX",
    "Bao_cao_toan_van_rut_gon_6_trang_giu_format_bang_landscape_tap_chi_ncv_app_5_muc_chinh.docx",
)


def set_run_font(run, size: float = 12, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    rpr.get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def replace_paragraph_text(paragraph, text: str, *, size: float = 12, italic: bool = False) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size=size, italic=italic)


def fitted_height(path: Path, width_cm: float) -> float:
    from PIL import Image

    with Image.open(path) as img:
        width_px, height_px = img.size
    return width_cm * height_px / width_px


def extract_image_from_paragraph(doc: Document, paragraph_index: int, out_path: Path) -> Path:
    paragraph = doc.paragraphs[paragraph_index]
    blip = next(paragraph._p.iter(qn("a:blip")), None)
    if blip is None:
        raise RuntimeError(f"Paragraph {paragraph_index} does not contain an image.")
    rid = blip.get(qn("r:embed"))
    part = doc.part.related_parts[rid]
    suffix = Path(str(part.partname)).suffix or ".png"
    path = out_path.with_suffix(suffix)
    path.write_bytes(part.blob)
    return path


def move_paragraph_before(doc: Document, paragraph, before_paragraph) -> None:
    body = doc.element.body
    body.remove(paragraph._p)
    index = list(body).index(before_paragraph._p)
    body.insert(index, paragraph._p)


def add_main_result_figure(doc: Document) -> None:
    discussion = next(p for p in doc.paragraphs if p.text.strip() == "4. BÀN LUẬN")
    with tempfile.TemporaryDirectory() as tmp:
        image_path = extract_image_from_paragraph(doc, 67, Path(tmp) / "boxplot")
        image_p = doc.add_paragraph()
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_p.paragraph_format.space_before = Pt(0)
        image_p.paragraph_format.space_after = Pt(1)
        width_cm = 12.8
        height_cm = min(fitted_height(image_path, width_cm), 5.5)
        image_p.add_run().add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))

        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.line_spacing = 1
        caption_p.paragraph_format.space_before = Pt(0)
        caption_p.paragraph_format.space_after = Pt(2)
        caption_text = (
            "Hình 3. Biểu đồ hộp phân bố góc vai và góc khuỷu theo từng video, "
            "hỗ trợ nhận diện mức dao động động tác giữa hai bài tập."
        )
        run = caption_p.add_run(caption_text)
        set_run_font(run, size=10, italic=True)

        move_paragraph_before(doc, caption_p, discussion)
        move_paragraph_before(doc, image_p, caption_p)


def main() -> None:
    doc = Document(str(SRC))

    replace_paragraph_text(
        doc.paragraphs[1],
        "Nhóm tác giả: Đinh Lê Quỳnh Phương, sinh viên CNCQKHDL1-1A; Kim Mạnh Hưng, sinh viên CNCQKHDL1-1A; Nguyễn Thị Thanh Nga, sinh viên CNCQKHDL1-1A; Nguyễn Thị Thơm, sinh viên CNCQ KTPHCN3-1A; Nguyễn Thị Thu Hương, sinh viên CNCQ YTCC22-1A - Trường Đại học Y tế Công cộng.",
    )
    replace_paragraph_text(
        doc.paragraphs[2],
        "Người báo cáo: Đinh Lê Quỳnh Phương, sinh viên CNCQKHDL1-1A, Trường Đại học Y tế Công cộng.",
    )
    replace_paragraph_text(
        doc.paragraphs[3],
        "Người hướng dẫn 1: TS. Trần Hồng Việt, giảng viên hướng dẫn, Viện Trí tuệ Nhân tạo, Trường Đại học Công nghệ, ĐHQG Hà Nội.",
    )
    replace_paragraph_text(
        doc.paragraphs[4],
        "Người hướng dẫn 2: CN. Nguyễn Thị Thùy Chi, giảng viên hướng dẫn, Trường Đại học Y tế Công cộng.",
    )

    intro = [
        "Phục hồi chức năng từ xa đang trở thành nhu cầu thiết yếu trong bối cảnh gánh nặng bệnh lý cơ xương khớp, chấn thương và đột quỵ ngày càng tăng. Trên phạm vi toàn cầu, nhu cầu phục hồi chức năng được ghi nhận ở khoảng một phần ba dân số, trong khi khả năng tiếp cận dịch vụ còn hạn chế tại nhiều quốc gia có nguồn lực y tế chưa đồng đều (1,2). Tại Việt Nam, thiếu hụt nhân lực phục hồi chức năng, khoảng cách địa lý và chi phí đi lại khiến nhiều người bệnh phải tự tập tại nhà sau xuất viện. Nếu không được giám sát, người bệnh dễ tập sai kỹ thuật, đau tăng, giảm tuân thủ điều trị và khó cung cấp dữ liệu khách quan cho bác sĩ/KTV khi tái khám (3).",
        "Trí tuệ nhân tạo và thị giác máy tính mở ra khả năng lượng hóa vận động bằng camera thông thường, không cần hệ thống đánh dấu chuyên dụng hoặc cảm biến đeo. Các tổng quan gần đây cho thấy mô hình ước lượng tư thế và đánh giá bài tập có thể hỗ trợ cá nhân hóa phản hồi, theo dõi tuân thủ và tạo dữ liệu định lượng cho nhân viên y tế trong môi trường tại nhà (4,7,8). MediaPipe Pose và OpenPose là hai hướng tiếp cận marker-less thường được sử dụng để trích xuất điểm mốc cơ thể, từ đó tính toán góc vận động, so sánh với ngưỡng tham chiếu hoặc đưa vào mô hình học máy để phân loại chất lượng động tác (5,6). Một số nghiên cứu quốc tế đã triển khai học sâu hoặc ứng dụng di động cho theo dõi bài tập phục hồi chức năng, nhưng kết quả phụ thuộc nhiều vào bối cảnh ghi hình, đặc điểm người bệnh và cách định nghĩa tiêu chuẩn đúng/sai (9-11).",
        "Tại Việt Nam, bằng chứng ứng dụng trên video người bệnh thực còn hạn chế. Môi trường ghi hình tại nhà hoặc phòng khám có thể xuất hiện che khuất tay, thay đổi góc quay, chuyển động bù trừ, ánh sáng không đồng đều và người hỗ trợ trong khung hình. Những yếu tố này có thể làm sai lệch điểm mốc cơ thể và ảnh hưởng đến độ chính xác của việc tính góc vai/khuỷu. Vì vậy, cần đánh giá tính khả thi của hệ thống trên dữ liệu thực tế trước khi mở rộng triển khai lâm sàng. Nghiên cứu này nhằm: (1) xây dựng mô hình nhận diện và đánh giá một số bài tập phục hồi chức năng khớp vai bằng thị giác máy tính; (2) so sánh kết quả của hệ thống với dữ liệu tham chiếu/chuyên gia trên bộ video người bệnh thực; và (3) mô tả khả năng tích hợp dashboard, video overlay và frame gallery trong quy trình giám sát từ xa.",
    ]
    for paragraph, text in zip(doc.paragraphs[7:10], intro):
        replace_paragraph_text(paragraph, text)

    add_main_result_figure(doc)
    doc.save(OUT)

    full_text = "\n".join(p.text for p in doc.paragraphs)
    main_text = full_text.split("PHỤ LỤC")[0]
    main_figs = sorted(set(re.findall(r"Hình\s+\d+", main_text)))
    main_tables = sorted(set(re.findall(r"Bảng\s+\d+", main_text)))
    print(OUT)
    print(f"main_figs={main_figs}")
    print(f"main_tables={main_tables}")
    print(f"main_total={len(main_figs) + len(main_tables)}")


if __name__ == "__main__":
    main()
