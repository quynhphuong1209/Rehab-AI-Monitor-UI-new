from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Bao_cao_toan_van_rut_gon_6_trang_giu_format_bang_landscape_tap_chi_ncv_app_5_muc_chinh.docx"
OUT = ROOT / "Bao_cao_toan_van_rut_gon_6_trang_giu_format_bang_landscape_tap_chi_ncv_app_5_muc_chinh_backend_frontend.docx"

APP_INTRO = [
    "Nguồn triển khai: https://rehab-ai-monitor.com/. Nguồn mã và README tham khảo: https://github.com/quynhphuong1209/Rehab-AI-Monitor-UI-new. Rehab AI Monitor được tổ chức theo kiến trúc web tách lớp, gồm giao diện người dùng ở phía frontend, lớp API/backend xử lý nghiệp vụ và lớp xử lý AI tạo dữ liệu phân tích từ video tập luyện.",
    "Về frontend, phiên bản web trong mã nguồn sử dụng React/Vite, trong đó web/src/App.tsx đảm nhiệm bố cục giao diện, điều hướng theo vai trò, màn hình đăng nhập, dashboard, bảng dữ liệu, video, frame gallery và phiếu đánh giá. web/src/api.ts là lớp client gọi API, còn web/src/styles.css định nghĩa hệ thống sidebar, thẻ chỉ số, tab, bảng, biểu đồ và responsive layout. Giao diện được thiết kế theo mô hình sidebar bên trái, thanh tiêu đề phía trên và vùng làm việc chính ở giữa, giúp người dùng truy cập chức năng đúng vai trò mà không phải chuyển qua nhiều hệ thống rời rạc.",
    "Về backend, hệ thống sử dụng FastAPI trong backend/main.py để cung cấp các endpoint chính như /auth/login, /dashboard, /videos/{identifier}/detail, /videos/upload, /videos/{identifier}/evaluations, nhóm analysis-jobs, exports, chart-exports, symptoms, schedules và admin/users. Backend chịu trách nhiệm xác thực tài khoản, phân quyền, lọc dữ liệu theo vai trò, đồng bộ dữ liệu từ các file JSON, cấp token truy cập media và trả payload thống nhất cho frontend hiển thị.",
    "Lớp dữ liệu runtime hiện được lưu và đồng bộ qua các tệp JSON trong database và thư mục gốc, gồm video_list.json, doctor_evaluations.json, patient_symptoms.json, research_data.json, lich_su_tap_luyen.json, schedules.json và latest_video_bundle.json. Cách tổ chức này cho phép dashboard đọc cùng một nguồn dữ liệu cho nhiều vai trò, đồng thời tách dữ liệu bệnh nhân, phiếu đánh giá, dữ liệu nghiên cứu, lịch nhắc và kết quả AI để thuận tiện kiểm tra, sao lưu và xuất báo cáo.",
    "Ở vai trò người bệnh, frontend ưu tiên các thao tác đơn giản: xem hướng dẫn/bài tập mẫu, khai báo triệu chứng và mức đau VAS, tải video tập luyện, theo dõi lịch nhắc và xem phản hồi sau đánh giá. Các thông tin trả về từ backend được trình bày dưới dạng nhận xét chuyên môn, kết quả AI, video/biểu đồ liên quan và khuyến nghị tập tiếp hoặc cần điều chỉnh kỹ thuật.",
    "Ở vai trò bác sĩ/KTV phục hồi chức năng, giao diện tập trung vào việc đối chiếu video gốc, video overlay, biểu đồ góc khớp và kết quả AI với quan sát lâm sàng. Phiếu đánh giá cho phép nhập kết quả, lỗi kỹ thuật, nhận xét gửi người bệnh, ghi chú cho nghiên cứu viên và chỉ định tiếp theo. Khi lưu, backend ghi nhận phiếu vào doctor_evaluations.json để tạo lớp dữ liệu chuyên môn dùng đối chiếu với kết quả AI.",
    "Ở vai trò nghiên cứu viên, màn hình Phân tích & trích xuất là vùng thao tác trung tâm. Người dùng có thể chọn video theo bệnh nhân/bài tập, cấu hình mô hình MediaPipe Heavy/Full/Lite, skip frame, resize width và confidence, theo dõi trạng thái job, mở video đã phân tích, xem khung xương MediaPipe, đối chiếu REF đúng/sai/gần đúng, train hoặc apply lớp ML, lưu/gửi video, frames và biểu đồ khi cần. Danh sách video hiển thị trạng thái đã phân tích, điểm AI và các nút truy cập nhanh Video, Biểu đồ, Khung xương.",
    "Dashboard nghiên cứu viên trình bày đồng thời lớp định lượng và lớp kiểm tra trực quan. Lớp định lượng gồm Accuracy, F1-score, MAE, ICC, số frame, PASS/NEAR/FAIL/UNKNOWN, biểu đồ góc khớp theo frame, phân bố kết quả, histogram, boxplot, radar chỉ số nghiên cứu và bảng chỉ số. Lớp trực quan nằm ở tab Video & Frames, trong đó video overlay hiển thị khung xương, góc khớp, nhãn REF/ML và trạng thái PASS/NEAR/FAIL theo thời gian; frame gallery có phân trang, bộ lọc PASS, NEAR, FAIL, UNKNOWN và Có ML, mỗi thẻ frame thể hiện ảnh overlay, góc vai/khuỷu, kết quả pose, model, REF YouTube, nhãn ML và xác suất ba lớp.",
    "Ở vai trò quản trị viên, frontend cung cấp không gian quản lý tài khoản, phân quyền, dữ liệu nghiên cứu, video đã phân tích, phiếu đánh giá và trạng thái vận hành. Backend áp dụng kiểm soát truy cập theo vai trò trước khi trả dữ liệu hoặc cho phép thao tác nhạy cảm như tạo tài khoản, khóa/mở tài khoản, chạy lại phân tích AI, lưu bundle video, xuất video/frames và quản lý lịch nhắc. Cách phân quyền này giúp bảo vệ dữ liệu người bệnh và giữ luồng nghiên cứu nhất quán từ thu thập video, xử lý AI, đánh giá chuyên môn đến phản hồi.",
    "Luồng AI sử dụng OpenCV để đọc video RGB theo frame, MediaPipe Pose/BlazePose để trích xuất 33 điểm mốc cơ thể, sau đó tính góc vai/khuỷu và so sánh với dữ liệu tham chiếu theo từng bài tập. Với Codman, hệ thống chia các giai đoạn và áp dụng ngưỡng sai số góc; với bài tập với gậy, kết quả được tổng hợp trên toàn video. Lớp RandomForestClassifier được dùng như mô hình ML tầng hai để hỗ trợ phân loại frame thành đúng, gần đúng, sai hoặc không xác định. Hai ví dụ minh họa trong phụ lục được chụp từ tài khoản nghiên cứu viên sau khi chọn video Hoàng Hạnh Nguyên - Codman và Nguyễn Thị Nga - bài tập với gậy, chờ dashboard và tab Video & Frames hiển thị đầy đủ; các ảnh thể hiện đồng thời dashboard định lượng và kiểm tra trực quan qua video/frame. Kết quả AI được sử dụng như công cụ hỗ trợ theo dõi và chuẩn hóa dữ liệu nghiên cứu, không thay thế kết luận lâm sàng của bác sĩ/KTV.",
]


def set_run_font(run) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(key), "Times New Roman")


def main() -> None:
    doc = Document(SRC)
    start = next(
        (i for i, paragraph in enumerate(doc.paragraphs) if "rehab-ai-monitor.com" in paragraph.text),
        None,
    )
    if start is None:
        raise RuntimeError("Cannot find the app introduction source paragraph.")

    for offset, text in enumerate(APP_INTRO):
        paragraph = doc.paragraphs[start + offset]
        paragraph.text = ""
        run = paragraph.add_run(text)
        set_run_font(run)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if offset == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
